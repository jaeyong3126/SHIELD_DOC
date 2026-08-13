from pathlib import Path
import re

from pypdf import PdfReader
from docx import Document


# TXT 파일에서 텍스트 추출
def parse_txt(file):
    if isinstance(file, (str, Path)):
        data = Path(file).read_bytes()

    else:
        file.seek(0)
        data = file.read()

    # 한글 TXT 파일 인코딩 대응
    for encoding in ["utf-8", "utf-8-sig", "cp949"]:
        try:
            return data.decode(encoding)

        except UnicodeDecodeError:
            pass

    raise ValueError("TXT 파일의 인코딩을 읽을 수 없습니다.")


# PDF 파싱 과정에서 생기는 불필요한 공백 정리
def clean_pdf_text(text):
    # 특수 공백 제거
    text = text.replace("\u00a0", " ")
    text = text.replace("\u200b", "")

    # 같은 줄 안의 연속 공백 정리
    text = re.sub(r"[ \t]{2,}", " ", text)

    # 이메일의 @, . 주변에 생긴 공백 보정
    email_pattern = re.compile(
        r"(?P<local>[A-Za-z0-9._%+-]+)"
        r"\s*@\s*"
        r"(?P<domain>"
        r"[A-Za-z0-9-]+"
        r"(?:\s*\.\s*[A-Za-z0-9-]+)+"
        r")"
    )

    def fix_email(match):
        local = match.group("local")

        domain = re.sub(
            r"\s*\.\s*",
            ".",
            match.group("domain")
        )

        return f"{local}@{domain}"

    return email_pattern.sub(
        fix_email,
        text
    )


# PDF 파일에서 텍스트 추출
def parse_pdf(file):
    if not isinstance(file, (str, Path)):
        file.seek(0)

    reader = PdfReader(file)
    pages = []

    for page in reader.pages:
        page_text = page.extract_text()

        if page_text:
            page_text = clean_pdf_text(
                page_text
            )

            pages.append(page_text)

    return "\n".join(pages)


# DOCX 파일에서 텍스트 추출
def parse_docx(file):
    if not isinstance(file, (str, Path)):
        file.seek(0)

    document = Document(file)
    texts = []

    # XML 내부의 실제 텍스트를 가져오는 함수
    def get_xml_text(element):
        text = "".join(
            node.text or ""
            for node in element.xpath(".//w:t")
        )

        return text.strip()

    # 일반 문단
    for paragraph in document.paragraphs:
        text = get_xml_text(
            paragraph._p
        )

        if text:
            texts.append(text)

    # 표
    for table in document.tables:
        for row in table.rows:

            row_text = []
            seen_cells = set()

            for cell in row.cells:

                # 병합 셀 중복 방지
                cell_id = id(cell._tc)

                if cell_id in seen_cells:
                    continue

                seen_cells.add(
                    cell_id
                )

                text = get_xml_text(
                    cell._tc
                )

                if text:
                    row_text.append(text)

            if row_text:
                texts.append(
                    " | ".join(row_text)
                )

    return "\n".join(texts)


# 파일 형식을 확인하고 텍스트 추출
def parse_document(file):
    # 로컬 파일 경로인 경우
    if isinstance(file, (str, Path)):
        filename = Path(file).name
        extension = Path(file).suffix.lower()

    # Streamlit UploadedFile 같은 파일 객체인 경우
    else:
        filename = file.name
        extension = Path(filename).suffix.lower()

    if extension == ".txt":
        text = parse_txt(file)

    elif extension == ".pdf":
        text = parse_pdf(file)

    elif extension == ".docx":
        text = parse_docx(file)

    else:
        raise ValueError(
            "지원하지 않는 파일 형식입니다. "
            "TXT, PDF, DOCX만 사용할 수 있습니다."
        )

    # 텍스트를 하나도 추출하지 못한 경우
    if not text.strip():
        raise ValueError(
            "문서에서 텍스트를 추출하지 못했습니다."
        )

    # 운영체제에 따라 다른 줄바꿈을 \n으로 통일
    text = text.replace(
        "\r\n",
        "\n"
    ).replace(
        "\r",
        "\n"
    )

    return {
        "filename": filename,
        "text": text
    }